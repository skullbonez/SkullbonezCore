"""
File: Agentic/Manuals/SkullbonezCoreManual/build_manual.py
Purpose:
  Builds the print-oriented Skullbonez Core technical manual as a DOCX.

Summary:
  Generates the manual's shared text, figures, DOCX layout, and PDF rendering
  from one source so both deliverables retain the same structure and references.

Mental model:
  The manual is generated because it contains repeated chapter furniture,
  figure plates, equation cards, and print geometry that are easier to keep
  consistent from code than by hand-editing Word paragraphs.

Glossary:
  DOCX: Office Open XML word-processing document used as the editable manual
    master.
  Render QA: Visual verification pass that renders document pages to inspect
    pagination, image placement, and text fit.

Invariants:
  - Output is documentation only; it must not modify engine source or runtime data.
  - Figures and equations are generated assets owned by the manual folder.
  - The DOCX remains the editable master; the PDF is produced by the render QA step.

Related:
  - Agentic/Manuals/SkullbonezCoreManual/Skullbonez_Core_Technical_Manual.docx
    is the generated editable artifact.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence
import textwrap
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table as RLTable,
    TableStyle as RLTableStyle,
)


ROOT = Path(__file__).resolve().parent
REPO_ROOT = ROOT.parents[2]
FIGURES = ROOT / "figures"
OUT_DOCX = ROOT / "Skullbonez_Core_Technical_Manual.docx"
OUT_PDF = REPO_ROOT / "output" / "pdf" / "Skullbonez_Core_Technical_Manual.pdf"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(26, 32, 44)
MUTED = RGBColor(94, 108, 132)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GRID = "CBD5E1"


@dataclass
class SectionSpec:
    title: str
    paragraphs: Sequence[str]
    bullets: Sequence[str] = ()
    figure: str | None = None
    equation: str | None = None
    table: tuple[Sequence[str], Sequence[Sequence[str]]] | None = None
    note: str | None = None


@dataclass
class ChapterSpec:
    number: int
    title: str
    thesis: str
    figure: str | None
    sections: Sequence[SectionSpec]
    improvements: Sequence[str]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    if bold:
        candidates = [
            Path("C:/Windows/Fonts/segoeuib.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
            Path("C:/Windows/Fonts/calibrib.ttf"),
        ] + candidates
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    text: str,
    xy: tuple[int, int],
    box_width: int,
    fill: str,
    fnt: ImageFont.FreeTypeFont,
    line_gap: int = 6,
    align: str = "left",
) -> int:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else current + " " + word
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= box_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)

    x, y = xy
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        line_width = bbox[2] - bbox[0]
        dx = 0
        if align == "center":
            dx = (box_width - line_width) // 2
        draw.text((x + dx, y), line, fill=fill, font=fnt)
        y += (bbox[3] - bbox[1]) + line_gap
    return y


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#4B5563") -> None:
    draw.line([start, end], fill=fill, width=4)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex >= sx else -1
        pts = [(ex, ey), (ex - sign * 14, ey - 8), (ex - sign * 14, ey + 8)]
    else:
        sign = 1 if ey >= sy else -1
        pts = [(ex, ey), (ex - 8, ey - sign * 14), (ex + 8, ey - sign * 14)]
    draw.polygon(pts, fill=fill)


def box(
    draw: ImageDraw.ImageDraw,
    xyxy: tuple[int, int, int, int],
    title: str,
    body: str = "",
    fill: str = "#F8FAFC",
    outline: str = "#94A3B8",
) -> None:
    x1, y1, x2, y2 = xyxy
    draw.rounded_rectangle(xyxy, radius=16, fill=fill, outline=outline, width=3)
    draw_wrapped(draw, title, (x1 + 18, y1 + 14), x2 - x1 - 36, "#0F172A", font(24, True), align="center")
    if body:
        draw_wrapped(draw, body, (x1 + 18, y1 + 60), x2 - x1 - 36, "#334155", font(17), line_gap=5, align="center")


def save_diagram(path: Path, title: str, painter) -> None:
    img = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 1050), fill="#FFFFFF")
    draw.text((60, 42), title, fill="#0B2545", font=font(42, True))
    draw.line((60, 104, 1740, 104), fill="#D8DEE9", width=3)
    painter(draw)
    img.save(path)


def make_equation(path: Path, title: str, lines: Sequence[str]) -> None:
    img = Image.new("RGB", (1600, 470), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((20, 20, 1580, 450), radius=22, fill="#FFFFFF", outline="#CBD5E1", width=3)
    draw.text((60, 48), title, fill="#0B2545", font=font(34, True))
    y = 120
    for line in lines:
        draw.text((90, y), line, fill="#111827", font=font(29))
        y += 54
    img.save(path)


def build_figures() -> dict[str, Path]:
    FIGURES.mkdir(parents=True, exist_ok=True)
    figures: dict[str, Path] = {}

    def add(name: str, title: str, painter) -> None:
        path = FIGURES / f"{name}.png"
        save_diagram(path, title, painter)
        figures[name] = path

    add(
        "architecture_map",
        "Whole-Engine Ownership Map",
        lambda d: (
            box(d, (720, 160, 1080, 285), "Run", "Composition root and runtime coordinator", "#EFF6FF"),
            box(d, (80, 390, 390, 555), "Runtime Shell", "Window, input, timers, config, capture, diagnostics", "#F8FAFC"),
            box(d, (450, 390, 760, 555), "Scene + Assets", "SceneRuntime, AssetSystem, TestScene, styles", "#F0FDF4"),
            box(d, (820, 390, 1130, 555), "Physics", "PhysicsEngine, PhysicsScene, PhysicsWorld, stores", "#FFF7ED"),
            box(d, (1190, 390, 1500, 555), "Rendering", "RuntimeRenderer, RenderGraph, DX12 backend", "#F5F3FF"),
            box(d, (930, 700, 1240, 865), "GameModelCollection", "Stable model order and model-owner sync boundary", "#FEFCE8"),
            box(d, (1320, 700, 1630, 865), "UI + Replay", "Frame data in, commands and scrub state out", "#FDF2F8"),
            arrow(d, (900, 285), (235, 390)),
            arrow(d, (900, 285), (605, 390)),
            arrow(d, (900, 285), (975, 390)),
            arrow(d, (900, 285), (1345, 390)),
            arrow(d, (975, 555), (1085, 700)),
            arrow(d, (1345, 555), (1475, 700)),
            arrow(d, (1085, 700), (1345, 555)),
        ),
    )

    add(
        "frame_loop",
        "Frame Loop and Fixed-Tick Dispatch",
        lambda d: (
            box(d, (80, 250, 305, 390), "Input", "Key edges, mouse, UI intent", "#F8FAFC"),
            box(d, (380, 250, 605, 390), "Commands", "Apply scene, replay, UI, capture intent", "#F0FDF4"),
            box(d, (680, 250, 905, 390), "SimulationSystem", "Accumulate wall time or fixed-step ticks", "#FFF7ED"),
            box(d, (980, 250, 1205, 390), "Physics", "Step one or more 1/120 s ticks", "#FFF1F2"),
            box(d, (1280, 250, 1505, 390), "Render", "Camera, passes, overlays, UI text", "#F5F3FF"),
            box(d, (670, 610, 930, 760), "Present", "Close command list, submit, flip, signal fence", "#EFF6FF"),
            arrow(d, (305, 320), (380, 320)),
            arrow(d, (605, 320), (680, 320)),
            arrow(d, (905, 320), (980, 320)),
            arrow(d, (1205, 320), (1280, 320)),
            arrow(d, (1390, 390), (820, 610)),
        ),
    )

    add(
        "scene_assets",
        "Scene and Asset Data Flow",
        lambda d: (
            box(d, (80, 190, 380, 340), "Scene JSON", "playback, simulation, objects, cameras, rendering", "#F8FAFC"),
            box(d, (460, 190, 760, 340), "Parser", "TestScene data and runtime overrides", "#F0FDF4"),
            box(d, (840, 190, 1140, 340), "AssetSystem", "Logical names to registered source assets", "#FEFCE8"),
            box(d, (1220, 190, 1520, 340), "Scene Setup", "Authored/generated setup builds runtime state", "#FFF7ED"),
            box(d, (290, 610, 590, 760), "Models", "GameModel order, bodies, colliders, materials", "#FFF1F2"),
            box(d, (710, 610, 1010, 760), "World", "Terrain, water, gravity, drag, cinematic overrides", "#EFF6FF"),
            box(d, (1130, 610, 1430, 760), "Render Inputs", "Cameras, styles, materials, pass resources", "#F5F3FF"),
            arrow(d, (380, 265), (460, 265)),
            arrow(d, (760, 265), (840, 265)),
            arrow(d, (1140, 265), (1220, 265)),
            arrow(d, (1370, 340), (440, 610)),
            arrow(d, (1370, 340), (860, 610)),
            arrow(d, (1370, 340), (1280, 610)),
        ),
    )

    add(
        "physics_pipeline",
        "Fixed Physics Tick Pipeline",
        lambda d: (
            box(d, (80, 160, 340, 280), "1. Force Pass", "Gravity, buoyancy, drag, tornado, pending impulses", "#FFF7ED"),
            box(d, (420, 160, 680, 280), "2. Broadphase", "SpatialGrid swept/static bounds produce candidates", "#EFF6FF"),
            box(d, (760, 160, 1020, 280), "3. Middle Phase", "Reachability, fixed/joint/sleep pruning", "#F8FAFC"),
            box(d, (1100, 160, 1360, 280), "4. CCD Front-End", "Advance to object time of impact only", "#FFF1F2"),
            box(d, (1440, 160, 1700, 280), "5. Terrain Detect", "Sweep terrain and append manifolds", "#F0FDF4"),
            box(d, (250, 550, 510, 670), "6. Row Build", "Object and terrain manifolds become rows", "#FEFCE8"),
            box(d, (590, 550, 850, 670), "7. PGS Solve", "Normal/friction impulses update velocities", "#F5F3FF"),
            box(d, (930, 550, 1190, 670), "8. Writeback", "Solved velocities and correction to stores", "#EFF6FF"),
            box(d, (1270, 550, 1530, 670), "9. Sleep + Integrate", "Joints, sleep islands, remaining pose time", "#F8FAFC"),
            arrow(d, (340, 220), (420, 220)),
            arrow(d, (680, 220), (760, 220)),
            arrow(d, (1020, 220), (1100, 220)),
            arrow(d, (1360, 220), (1440, 220)),
            arrow(d, (1570, 280), (380, 550)),
            arrow(d, (510, 610), (590, 610)),
            arrow(d, (850, 610), (930, 610)),
            arrow(d, (1190, 610), (1270, 610)),
        ),
    )

    def contact_row(d: ImageDraw.ImageDraw) -> None:
        d.ellipse((260, 360, 560, 660), fill="#DBEAFE", outline="#2563EB", width=5)
        d.ellipse((960, 330, 1300, 670), fill="#FEF3C7", outline="#D97706", width=5)
        d.text((355, 485), "Body A", fill="#0F172A", font=font(33, True))
        d.text((1068, 485), "Body B", fill="#0F172A", font=font(33, True))
        d.ellipse((760, 495, 800, 535), fill="#DC2626")
        d.text((725, 550), "contact point", fill="#991B1B", font=font(24, True))
        arrow(d, (780, 515), (780, 300), "#DC2626")
        d.text((810, 340), "normal n", fill="#991B1B", font=font(26, True))
        arrow(d, (780, 515), (660, 445), "#475569")
        arrow(d, (780, 515), (910, 455), "#475569")
        d.text((625, 405), "rA", fill="#334155", font=font(24, True))
        d.text((920, 420), "rB", fill="#334155", font=font(24, True))
        d.line((650, 735, 950, 735), fill="#059669", width=5)
        d.line((800, 590, 800, 880), fill="#059669", width=5)
        d.text((980, 716), "two tangent axes span friction", fill="#065F46", font=font(26, True))

    add("contact_row", "Anatomy of a Persistent Contact Row", contact_row)

    add(
        "solver_loop",
        "Persistent Contact Solver Loop",
        lambda d: (
            box(d, (95, 245, 365, 405), "Manifold Rows", "bodyA/bodyB, point, normal, tangents, penetration", "#FEFCE8"),
            box(d, (460, 245, 730, 405), "Precompute", "effective mass, bias, friction budget, cache lookup", "#EFF6FF"),
            box(d, (825, 245, 1095, 405), "Warm Start", "apply cached lambda to scratch velocities", "#F0FDF4"),
            box(d, (1190, 245, 1460, 405), "PGS Iterations", "clamp normal and tangent impulses", "#F5F3FF"),
            box(d, (645, 650, 915, 810), "Writeback + Cache", "store solved velocities and next-frame impulses", "#FFF7ED"),
            arrow(d, (365, 325), (460, 325)),
            arrow(d, (730, 325), (825, 325)),
            arrow(d, (1095, 325), (1190, 325)),
            arrow(d, (1325, 405), (780, 650)),
            arrow(d, (780, 650), (600, 405)),
        ),
    )

    add(
        "render_pipeline",
        "Runtime Rendering Pass Order",
        lambda d: (
            box(d, (75, 185, 300, 310), "Shadow Maps", "terrain and object depth", "#F8FAFC"),
            box(d, (360, 185, 585, 310), "Reflection", "water texture, DXR or mirrored raster", "#EFF6FF"),
            box(d, (645, 185, 870, 310), "Scene Target", "HDR target when cinematic mode wraps frame", "#F5F3FF"),
            box(d, (930, 185, 1155, 310), "Objects", "body streams, materials, focus masks", "#FFF7ED"),
            box(d, (1215, 185, 1440, 310), "Terrain", "receiver, relief, shadows", "#F0FDF4"),
            box(d, (1500, 185, 1725, 310), "Water", "reflection and waves", "#EFF6FF"),
            box(d, (360, 600, 585, 725), "Debug", "broadphase, contacts, physics overlays", "#FEFCE8"),
            box(d, (645, 600, 870, 725), "Volumetric", "half-res shafts and fog input", "#F5F3FF"),
            box(d, (930, 600, 1155, 725), "Tonemap", "HDR resolve, bloom, grade", "#FFF7ED"),
            box(d, (1215, 600, 1440, 725), "UI/Text", "HUD, panels, replay controls", "#F8FAFC"),
            arrow(d, (300, 247), (360, 247)),
            arrow(d, (585, 247), (645, 247)),
            arrow(d, (870, 247), (930, 247)),
            arrow(d, (1155, 247), (1215, 247)),
            arrow(d, (1440, 247), (1500, 247)),
            arrow(d, (1615, 310), (475, 600)),
            arrow(d, (585, 662), (645, 662)),
            arrow(d, (870, 662), (930, 662)),
            arrow(d, (1155, 662), (1215, 662)),
        ),
    )

    add(
        "render_graph",
        "Render Graph Resource State Tracking",
        lambda d: (
            box(d, (100, 190, 410, 330), "Pass Declarations", "reads[], writes[], resource handles", "#F8FAFC"),
            box(d, (520, 190, 830, 330), "Compile", "walk passes and track current access", "#EFF6FF"),
            box(d, (940, 190, 1250, 330), "Transitions", "emit before -> after records when access changes", "#FEFCE8"),
            box(d, (1360, 190, 1670, 330), "DX12 Barriers", "translate graph access to concrete states", "#FFF7ED"),
            box(d, (520, 610, 830, 760), "Transient Plan", "first/last pass, pool slots, descriptor high water", "#F5F3FF"),
            box(d, (940, 610, 1250, 760), "Materialize", "physical textures, RTV/DSV/SRV/UAV rows", "#F0FDF4"),
            arrow(d, (410, 260), (520, 260)),
            arrow(d, (830, 260), (940, 260)),
            arrow(d, (1250, 260), (1360, 260)),
            arrow(d, (675, 330), (675, 610)),
            arrow(d, (830, 685), (940, 685)),
        ),
    )

    add(
        "dx12_memory",
        "DX12 Frame Ownership and Lifetime",
        lambda d: (
            box(d, (100, 195, 390, 345), "CPU Frame", "command list, allocator, upload arena", "#EFF6FF"),
            box(d, (500, 195, 790, 345), "Descriptors", "persistent source rows copied to transient visible heap", "#F8FAFC"),
            box(d, (900, 195, 1190, 345), "GPU Work", "execute command list asynchronously", "#F5F3FF"),
            box(d, (1300, 195, 1590, 345), "Fence", "signals when allocator/upload/resources are reusable", "#FEFCE8"),
            box(d, (500, 635, 790, 785), "Deferred Release", "resources retire after completed fence", "#FFF7ED"),
            box(d, (900, 635, 1190, 785), "Readback", "timers and captures wait on scoped fences", "#F0FDF4"),
            arrow(d, (390, 270), (500, 270)),
            arrow(d, (790, 270), (900, 270)),
            arrow(d, (1190, 270), (1300, 270)),
            arrow(d, (1445, 345), (1045, 635)),
            arrow(d, (1445, 345), (645, 635)),
        ),
    )

    add(
        "validation_map",
        "Validation and Evidence Flow",
        lambda d: (
            box(d, (100, 210, 390, 360), "Change Area", "docs, physics, DX12, perf, scenes, tools", "#F8FAFC"),
            box(d, (500, 210, 790, 360), "Targeted Gate", "smallest script that proves the risk", "#EFF6FF"),
            box(d, (900, 210, 1190, 360), "Artifacts", "CSV baselines, screenshots, logs, SkullScope queries", "#FEFCE8"),
            box(d, (1300, 210, 1590, 360), "Handoff", "command, result, log path, known residual risk", "#F0FDF4"),
            box(d, (500, 620, 790, 770), "Do Not Validate", "documentation-only prose and print artifacts", "#FFF7ED"),
            box(d, (900, 620, 1190, 770), "Never Claim", "no validation success without output", "#FFF1F2"),
            arrow(d, (390, 285), (500, 285)),
            arrow(d, (790, 285), (900, 285)),
            arrow(d, (1190, 285), (1300, 285)),
            arrow(d, (645, 360), (645, 620)),
            arrow(d, (1045, 360), (1045, 620)),
        ),
    )

    equations = {
        "eq_fixed_step": (
            "Fixed timestep accumulator",
            [
                "acc = acc + frame_dt * time_scale",
                "while acc >= h: Step(h); acc = acc - h",
                "Skullbonez physics tick: h = 1 / 120 seconds",
            ],
        ),
        "eq_euler": (
            "Velocity-first integration",
            [
                "v_next = v + h * M_inverse * F",
                "x_next = x + h * v_next",
                "orientation advances from angular velocity after solver velocity writeback",
            ],
        ),
        "eq_contact_velocity": (
            "Contact point velocity",
            [
                "vA_p = vA + omegaA cross rA",
                "vB_p = vB + omegaB cross rB",
                "v_rel = vB_p - vA_p,    v_n = dot(v_rel, n)",
            ],
        ),
        "eq_effective_mass": (
            "Normal row impulse update",
            [
                "m_eff = 1 / (invMassA + invMassB + angular terms)",
                "delta_lambda_n = m_eff * (bias - v_n)",
                "lambda_n = max(0, lambda_n + delta_lambda_n)",
            ],
        ),
        "eq_friction": (
            "Tangent friction clamp",
            [
                "lambda_t = (lambda_t1, lambda_t2)",
                "length(lambda_t) <= friction_limit",
                "Skullbonez clamps the tangent pair as one 2D cone budget.",
            ],
        ),
        "eq_render_transform": (
            "World to clip space",
            [
                "p_world = M_model * p_local",
                "p_clip = Projection * View * p_world",
                "Shadow/reflection passes use alternate View matrices, not alternate objects.",
            ],
        ),
        "eq_graph_transition": (
            "Render graph transition rule",
            [
                "before = tracked_access(resource)",
                "if before != requested_access: emit transition(before -> requested)",
                "tracked_access(resource) = requested_access",
            ],
        ),
        "eq_fence": (
            "Fence-scoped reuse",
            [
                "submit(command_list); signal(frame_fence)",
                "reuse allocator/upload/descriptors only after completed_fence >= frame_fence",
                "deferred releases retire on completed fences.",
            ],
        ),
    }
    for name, (title, lines) in equations.items():
        path = FIGURES / f"{name}.png"
        make_equation(path, title, lines)
        figures[name] = path

    return figures


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = GRID, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths: Sequence[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Skullbonez Core Technical Manual"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.add_run("Page ")
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED

    props = doc.core_properties
    props.title = "Skullbonez Core Technical Manual"
    props.subject = "Engine architecture, physics, rendering, diagnostics, and improvement notes"
    props.author = "SkullbonezCore"
    props.keywords = "SkullbonezCore, engine, physics, rendering, DX12, manual"


def add_para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead) :])
    else:
        p.add_run(text)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.35])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_borders(cell, "D8DEE9", "4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], LIGHT_BLUE)
        set_cell_borders(hdr[i])
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_borders(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def chapter_specs() -> list[ChapterSpec]:
    return [
        ChapterSpec(
            1,
            "Engine Architecture",
            "Skullbonez Core is easiest to understand as a runtime composition root wrapped around explicit subsystem boundaries. Run still coordinates the frame, but newer facades name the work that is being carved out of it.",
            "architecture_map",
            [
                SectionSpec(
                    "The Composition Root",
                    [
                        "Run owns the long-lived runtime graph: native window access, config, scene queue, cameras, world environment, game models, UI, replay, capture, diagnostics, and the renderer host. It is still broad, but the code now labels its borrowed boundaries so future extraction work has real names.",
                        "The key habit is to read ownership before reading behavior. EngineContext is a bound view over Run-owned systems. RuntimeRenderModelFrameView publishes scene-owned render values for one call, while RuntimeRenderer keeps its concrete resource owners. PhysicsModelAccess is the model-owner sync boundary between legacy GameModel storage and physics stores.",
                        "That pattern lets the engine modernize without pretending the old global shape vanished overnight. A facade is useful only when it tells the truth about ownership and lifetime.",
                    ],
                    bullets=[
                        "Run is the composition root, not a domain model.",
                        "SceneRuntime owns active scene and queue state, while Run still applies heavy load/reset side effects.",
                        "SimulationSystem owns timestep policy and accumulators.",
                        "RuntimeRenderer owns pass objects and the ordered render frame.",
                    ],
                    note="Reading rule: before changing a subsystem, identify the owner, the borrowed view, and the writeback point.",
                ),
                SectionSpec(
                    "The Frame as a Contract",
                    [
                        "A frame starts with input and runtime commands, not with rendering. InputController captures key and mouse edges; UI emits command intent; scene and replay code may change the active state before simulation asks for time.",
                        "SimulationSystem then converts wall time or fixed-step mode into zero or more physics ticks. Camera and UI use a frame-level delta, but physics advances in fixed units so solver behavior stays reproducible.",
                        "Rendering follows after simulation and replay preview state have been applied. Run sets the active camera, publishes the model frame view, calls RuntimeRenderer::RenderFrameEntry, then restores any replay pose that was temporarily borrowed for presentation.",
                    ],
                    figure="frame_loop",
                    equation="eq_fixed_step",
                ),
                SectionSpec(
                    "Stable Indices and Data Boundaries",
                    [
                        "GameModelCollection is the stable owner of model order. Physics stores, render streams, debug overlays, replay samples, and scene snapshots all depend on that deterministic ordering.",
                        "The current engine still mirrors solved physics back into GameModel because editor, replay, render, and diagnostics consumers read model state. The important architectural progress is that the solver itself runs over PhysicsBodyStore, ColliderStore, compact streams, and bounded side-effect arrays.",
                        "This is why migration cleanup focuses so much on authority. A hot loop should not reach through Run or GameModel; it should consume compact value records and let the owner apply side effects after the pass.",
                    ],
                    bullets=[
                        "Authoritative legacy storage: GameModel vector inside GameModelCollection.",
                        "Physics hot path: aligned body-field arrays, collider records, body streams, solver scratch, side-effect buffers.",
                        "Render path: render instance streams and per-frame RuntimeRenderModelFrameView values.",
                        "Diagnostics path: bounded pipeline records and queryable SkullScope traces.",
                    ],
                ),
                SectionSpec(
                    "Subsystem Extraction Philosophy",
                    [
                        "The codebase is not doing a big-bang rewrite. Instead, it names a boundary, preserves behavior, proves the validation surface, and then tightens authority in later slices.",
                        "Good extraction removes hidden reach-through. Bad extraction merely gives the old reach-through a new name. The repository rules deliberately discourage vague compatibility nouns because a bridge without a deletion condition can become permanent architecture.",
                        "This matters most in physics and rendering. Both are hot-path systems where callback chains, polymorphic services, and owner-side commands can silently turn a deterministic loop into scattered state mutation.",
                    ],
                    table=(
                        ["Pattern", "Useful When", "Risk"],
                        [
                            ["Facade", "It names a real boundary and forwards without reordering.", "It can hide old ownership if the next owner is not clear."],
                            ["Value record", "Hot path needs predictable memory and deterministic iteration.", "It must be refreshed and invalidated at the right times."],
                            ["Side-effect queue", "Solver or worker pass needs to report owner mutations later.", "It needs bounded capacity and deterministic commit order."],
                        ],
                    ),
                ),
                SectionSpec(
                    "The Mental Stack",
                    [
                        "At the top is runtime intent: command-line flags, scene JSON, UI commands, replay state, and validation harnesses. In the middle are owner subsystems that translate intent into current state. At the bottom are hot loops that should see only the data they need.",
                        "Most engine questions become easier when you place them on that stack. A scene option belongs near parsing and setup. A DX12 resource state belongs in graph/backend code. A contact impulse belongs in solver scratch until the writeback stage.",
                        "The manual follows that stack throughout: ownership first, data movement second, equations third, improvement notes last.",
                    ],
                    bullets=[
                        "Intent: command-line, scene JSON, UI, replay, validation.",
                        "Owners: Run, SceneRuntime, SimulationSystem, PhysicsScene, RuntimeRenderer, AssetSystem.",
                        "Hot data: aligned body-field arrays, collider records, render instances, descriptor rows, upload arenas.",
                    ],
                ),
            ],
            [
                "Continue shrinking Run by moving load/reset side effects behind named scene/runtime coordinators.",
                "Finish strict authority work so physics and rendering consume stores/streams without legacy model reach-through.",
                "Keep facades honest with owner, reason, deletion condition, and checker budget when they are transitional.",
                "Add more human-facing diagrams beside the class-structure reference as subsystems stabilize.",
            ],
        ),
        ChapterSpec(
            2,
            "Scene, Assets, and Runtime State",
            "Scenes are declarative inputs. They become runtime state only through parser, asset, setup, and model-owner boundaries.",
            "scene_assets",
            [
                SectionSpec(
                    "Scene Files as Runtime Requests",
                    [
                        "A Skullbonez scene file is a request, not a live object graph. It can describe playback, simulation, world overrides, cameras, objects, terrain, rendering, UI state, and validation gates.",
                        "SceneRuntime and SceneController hold the active queue/index state, while authored and generated setup helpers build the runtime objects. This keeps scene selection separate from the actual mutation that creates terrain, models, cameras, and validation records.",
                        "The runtime also accepts suites, generated demo modes, hero scene aliases, replay probes, live style control, and validation-only launch flags. All of those paths eventually have to converge on the same state ownership rules.",
                    ],
                    bullets=[
                        "Scene JSON is versioned declarative data.",
                        "Suites sequence scene files for validation and render capture.",
                        "Style files layer cinematic/material settings without rebuilding physics.",
                        "Snapshot writing serializes current state back to scene JSON when requested.",
                    ],
                ),
                SectionSpec(
                    "Asset Resolution",
                    [
                        "AssetSystem owns source asset identities and resolves logical names to data-root paths. Textures, shader programs, style files, scenes, asset libraries, terrain, and fonts all become source records before runtime code asks the backend to create GPU objects.",
                        "The asset registry is intentionally separate from backend memory. A texture source asset says what to load and how; the DX12 texture entry, descriptors, and resource lifetime are renderer concerns.",
                        "Reusable placeables should be registered asset libraries. Scenes should prefer assetInstances for reusable props, hulls, and compound structures instead of baking every part into scene object lists.",
                    ],
                    table=(
                        ["Asset Kind", "Runtime Use", "Owner Boundary"],
                        [
                            ["Texture2D", "Materials, sky, UI, water, terrain.", "AssetSystem source record; renderer GPU lifetime."],
                            ["ShaderProgram", "Logical shader names and contracts.", "AssetSystem resolves; ShaderDX12 compiles and reflects."],
                            ["AssetLibrary", "Reusable convex hulls and compounds.", "Registered from Run setup, expanded by scene setup."],
                        ],
                    ),
                ),
                SectionSpec(
                    "GameModelCollection as Traffic Controller",
                    [
                        "GameModelCollection is where authored objects become a deterministic order that physics, rendering, replay, and diagnostics can agree on. It is not merely a container; it is the model-owner boundary.",
                        "PhysicsModelAccess exposes the operations physics is allowed to perform: reload bodies, refresh colliders, write back one body or all bodies, access streams, and invalidate stream caches after solver mutation.",
                        "Render code does not need the entire collection as an owner; it needs per-frame views and render instances. That is why GameModelStreamProvider and RuntimeRenderModelFrameView matter: they mark the path from historical objects toward view-shaped data.",
                    ],
                    bullets=[
                        "Preserve model insertion order for deterministic solver and replay behavior.",
                        "Refresh collider snapshots when topology changes, not inside every row solve.",
                        "Invalidate derived streams after physics writeback.",
                        "Keep editor/replay mirrors outside the narrow solver loop.",
                    ],
                ),
                SectionSpec(
                    "World State",
                    [
                        "WorldEnvironment names gravity, fluid, drag, water height, and related runtime forces. Physics receives a world-force snapshot so the solver does not reach back into world ownership during the hot pass.",
                        "Terrain has two jobs: render geometry and collision surface. The physics path treats terrain contact as a borrowed per-call view and produces TerrainContactManifold value reports rather than mutating a full scene object.",
                        "Cinematic overrides change render presentation, not physics terrain. This separation is important: visual relief, fog, bloom, sky atmosphere, and water presentation must not secretly alter deterministic physics baselines.",
                    ],
                    equation="eq_euler",
                    note="Good boundary test: a render-only scene style should not reset physics, rebuild bodies, or change the physics regression CSV.",
                ),
                SectionSpec(
                    "Runtime Control Surfaces",
                    [
                        "Skullbonez has several control surfaces: command-line flags, scene directives, UI tabs, replay scrub controls, style harness files, and validation scripts. Their common task is to express intent without breaking subsystem ownership.",
                        "For example, the live style harness applies style-only JSON without reloading physics. Replay scrub temporarily applies presentation or solver samples for rendering, then restores live state after the frame. Validation launches use explicit flags so artifacts can be reproduced.",
                        "The more control surfaces a runtime has, the more valuable small owner-shaped structs become. They make each mutation explainable and reduce the chance that a convenience path bypasses the normal gate.",
                    ],
                    bullets=[
                        "Command-line flags are public compatibility surface.",
                        "UI emits commands; Run/runtime controllers apply them.",
                        "Replay preview borrows render state and restores it.",
                        "Validation scenes encode evidence needs in data.",
                    ],
                ),
            ],
            [
                "Move more scene reset/load side effects out of Run into scene-owned coordinators.",
                "Keep asset library registration data-driven and avoid new editor-only hardcoded compound objects.",
                "Broaden source diagrams for scene parsing, style layering, and snapshot writing.",
                "Add targeted docs for assetInstance authoring, hull baking, and reusable prop validation.",
            ],
        ),
        ChapterSpec(
            3,
            "Physics",
            "Physics is a fixed-step, deterministic data pipeline. Forces update velocities first; contact rows then solve impulses directly into velocities; final pose integration uses those solved velocities.",
            "physics_pipeline",
            [
                SectionSpec(
                    "Fixed Tick Dispatch",
                    [
                        "SimulationSystem owns the policy that turns frame time into physics steps. In deterministic fixed-step mode it ignores wall-clock accumulation and commits whole PHYSICS_FIXED_DT ticks from the time-scale accumulator. In variable-time scenes it still runs physics in fixed-size steps, capped to avoid runaway catch-up.",
                        "The important split is that physics receives a fixed dt while camera/UI presentation can use a frame-level dt. This is why scene validation can reproduce byte-exact physics behavior even when rendering cadence changes.",
                        "Every fixed tick runs through GameModelCollection::RunPhysics, PhysicsEngine::Step, PhysicsScene::RunPhysics, and PhysicsWorld::RunPhysics. PhysicsScene reloads cold metadata and initial hot fields, refreshes colliders when topology changed, runs the world over aligned hot arrays, copies sleep state, and writes the model-owner mirror once.",
                    ],
                    equation="eq_fixed_step",
                    table=(
                        ["Layer", "Responsibility"],
                        [
                            ["SimulationSystem", "Accumulates time and decides how many fixed ticks to commit."],
                            ["PhysicsEngine", "Runtime-facing facade over PhysicsScene."],
                            ["PhysicsScene", "Owns body/collider stores and model-owner synchronization."],
                            ["PhysicsWorld", "Runs broadphase, narrowphase, solver, sleep, diagnostics."],
                        ],
                    ),
                ),
                SectionSpec(
                    "Force Pass",
                    [
                        "The first major physics phase applies world forces to awake dynamic bodies. Fixed bodies return immediately. Sleeping bodies keep cached state and have their remaining time cleared until a contact or scene change wakes them.",
                        "Forces include gravity, buoyancy, drag, tornado effects, and pending impulses. These forces update linear and angular velocity. They do not resolve contacts.",
                        "This ordering matters because contact rows solve against the velocities that already include world forces. Gravity makes a body want to fall; the contact solver then produces impulses that oppose that motion when a support row exists.",
                    ],
                    equation="eq_euler",
                    bullets=[
                        "Sleeping bodies are static anchors until wake policy admits them.",
                        "A newly woken body applies forces immediately so it behaves like it was awake for the tick.",
                        "Tornado and world-force effects run before broadphase and row solving.",
                    ],
                ),
                SectionSpec(
                    "Broadphase and Middle Phase",
                    [
                        "SpatialGrid is the cheap candidate generator. It inserts swept bounds for fast moving dynamic bodies and static bounds for ordinary bodies, then emits unique pairs from occupied cells.",
                        "Candidate pairs are not contacts. They are only pairs worth testing more precisely. The middle phase prunes pairs that cannot affect simulation: swept bounding-radius reachability rejects, fixed-fixed pairs, point-joint pairs, and sleep-sleep pairs.",
                        "The grid is still populated with sleeping bodies because an awake body must be able to find and wake a sleeping neighbor. What gets removed is work that cannot produce wake, response, or diagnostics.",
                    ],
                    bullets=[
                        "Cell size follows the largest active broadphase primitive, clamped by config.",
                        "Fast-small sweep augmentation protects tiny high-speed projectiles.",
                        "Pair ordering is deterministic: pairs are normalized so A < B.",
                        "The middle phase may keep false positives, but it must not reject real possible contacts.",
                    ],
                ),
                SectionSpec(
                    "Object CCD Front-End",
                    [
                        "The object/object swept front-end is continuous collision detection for timing and wake-up. It can refine a time of impact and advance bodies to the contact candidate, but it does not apply object/object collision response.",
                        "That distinction is the core answer to the question that inspired this manual: contact rows do not become forces that later pass through F = ma. The world force pass already happened. Contact rows become impulses applied directly to solver velocities.",
                        "If a pair is already persistent and motion is small, CCD may skip the swept front-end and let the persistent manifold handle the settled contact. New or fast contacts keep the swept path so tunneling is avoided.",
                    ],
                    figure="contact_row",
                    equation="eq_contact_velocity",
                    note="Short version: CCD finds when to meet; persistent rows decide how the bodies respond.",
                ),
                SectionSpec(
                    "Terrain Detection and Manifolds",
                    [
                        "Terrain is still detected with a swept path before the shared solve, because fast bodies need a correct time of impact. But the terrain phase now emits terrain manifolds rather than running a separate terrain impulse response.",
                        "A TerrainContactManifold records bodyA, bodyB = -1 for terrain, normal, tangents, contact points, time of impact, support policy, friction permission, sleep inhibition, and feature ids.",
                        "Stable resting terrain support may seed sleep and warm starting. Edge or point contacts can inhibit sleep and avoid rest-only policy while still producing rows for normal and friction response.",
                    ],
                    bullets=[
                        "Terrain body index -1 means infinite mass, zero velocity, and no writeback.",
                        "Terrain normals are flipped when converted to shared PersistentContact row convention.",
                        "Terrain warm-start seed is strongest for stable support and smaller for sleep-inhibiting edge contacts.",
                    ],
                ),
                SectionSpec(
                    "Manifold to Row Build",
                    [
                        "PersistentContactSolver builds object/object rows from exact shape-pair manifolds and appends terrain rows from terrain manifolds. A row carries body indices, contact normal, tangent axes, rA/rB arms, penetration, feature id, support flags, friction policy, cached impulses, and per-row masses/biases.",
                        "Object rows use exact sphere, box, and convex-hull manifold builders. Quiet same-shape face footprints can reduce many manifold points to two well-spread cached rows, which cuts solver work without losing the support plane.",
                        "Each row gets a stable key. For object/object contacts the key combines the ordered pair and feature id. For terrain rows it uses a terrain kind bit, body A, and the terrain feature id.",
                    ],
                    figure="solver_loop",
                    equation="eq_effective_mass",
                ),
                SectionSpec(
                    "Projected Gauss-Seidel",
                    [
                        "The row solver is Catto-style sequential impulses. Precompute builds tangent axes, normal/tangent effective masses, restitution or Baumgarte bias, friction limits, and cached warm-start impulses.",
                        "Warm starting applies cached lambda to scratch solver velocities before iteration. Then each Projected Gauss-Seidel pass recomputes contact relative velocity, solves a normal impulse, clamps accumulated normal lambda to be non-negative, solves tangent friction, clamps the two tangents as one 2D friction vector, and applies only the delta impulse.",
                        "The solver writes velocity changes to compact SolverBodyState scratch first. Only after the row loop are solved linear and angular velocities written back to the authoritative hot-field arrays.",
                    ],
                    equation="eq_friction",
                    bullets=[
                        "Normal impulses can push apart but cannot glue together.",
                        "Friction is bounded by row policy: terrain can use normal impulse or warm-start support; object rows may use a gravity-sized budget or normal-coupled friction.",
                        "A deterministic early-out can stop iterations when total squared impulse delta is tiny.",
                    ],
                ),
                SectionSpec(
                    "Writeback, Integration, and Sleep",
                    [
                        "After solving, terrain rest policy can damp tiny supported rolling motion. Then solved velocities are written back to the authoritative hot arrays and queued for model-owner mirror writeback. Partial positional correction removes visible leftover overlap without replacing the velocity solver.",
                        "PhysicsWorld then wakes point-joint-connected sleepers, solves point joints, appends joint support edges, propagates sleep support, and integrates remaining time for awake models using the solved velocities.",
                        "Sleep is island-level. The engine builds contact and joint islands, checks quiet linear/angular thresholds, requires credible support anchors, honors terrain inhibition and point-joint error gates, then sleeps the whole eligible island only after every awake member accumulates the required quiet frames.",
                    ],
                    bullets=[
                        "Solved velocities are cached for next-frame warm starting.",
                        "Remaining pose integration uses each body's remaining time after CCD/terrain advancement.",
                        "Sleeping transition zeroes residual velocities to prevent tiny drift from reappearing on wake.",
                    ],
                ),
            ],
            [
                "Finish strict store authority so solver-facing code no longer needs legacy GameModel mirror assumptions.",
                "Continue replacing hot-path callbacks and compatibility commands with compact side-effect arrays.",
                "Consider broader awake/sleeping broadphase partitioning only with physics and perf validation evidence.",
                "Expand SkullScope canned questions for contact rows, sleep islands, and terrain support classification.",
                "Add diagrams directly beside the in-game physics pipeline overlay modes.",
            ],
        ),
        ChapterSpec(
            4,
            "Rendering",
            "Rendering is a DX12-first pass pipeline wrapped by RuntimeRenderer. The render graph is the vocabulary for resource access; the DX12 backend turns that vocabulary into explicit barriers, descriptors, uploads, fences, and presents.",
            "render_pipeline",
            [
                SectionSpec(
                    "RuntimeRenderer and Pass Ownership",
                    [
                        "RuntimeRenderer owns pass instances and the stable pass order: full-screen helpers, sky, scene target, shadows, reflection, objects, terrain, water, tornado visuals, debug overlays, volumetric light, tonemap, and UI text.",
                        "Run prepares the active camera and replay render state, checks that the process-bound backend is ready, narrows the backend into command/resource/diagnostic/raytracing facets, and calls RuntimeRenderer::RenderFrame with per-frame borrowed inputs.",
                        "Passes should not store broad runtime state. RenderFrameContext and RenderResourceContext separate per-frame drawing data from create/rebuild-only resource access.",
                    ],
                    bullets=[
                        "RenderFrameContext carries matrices, light vector, water height, scene view, backend facets, and dimensions.",
                        "RenderResourceContext is for resize-sensitive resources and lazy creation.",
                        "UI text may use the optional DXR facet when the active backend publishes it.",
                    ],
                    equation="eq_render_transform",
                ),
                SectionSpec(
                    "DX12 as the Production Backend",
                    [
                        "DX12 is the only runtime renderer. The historical GL and DX11 parity evidence is archived; renderer regression evidence now means DX12 validation errors plus DX12 screenshot/baseline comparisons.",
                        "RenderBackendDX12 implements the engine-facing renderer facade and owns explicit DX12 machinery: device/queue/swapchain/fence timeline, descriptor allocators, upload arenas, readback buffers, command list state, pipeline state cache, textures, meshes, framebuffers, and optional DXR structures.",
                        "The backend is deliberately explicit because DX12 makes resource state, descriptor lifetime, upload memory, and CPU/GPU synchronization visible. Hidden state is where GPU races like to hide.",
                    ],
                    table=(
                        ["DX12 Piece", "What It Owns"],
                        [
                            ["Dx12RenderDevice", "Factory/device/queue/swap chain/fence lifetime."],
                            ["Descriptor allocators", "SRV/UAV shader-visible rows plus CPU RTV/DSV rows."],
                            ["Frame upload system", "Per-frame upload arena and allocator pacing."],
                            ["Readback buffers", "Backbuffer capture and GPU timer readback."],
                        ],
                    ),
                ),
                SectionSpec(
                    "Render Graph Access",
                    [
                        "RenderGraph is the API-neutral contract for resources, passes, access intent, callback scheduling, and transient lifetimes. The current compiler is deliberately simple: start at each resource initial access, walk passes in added order, process reads and writes, record advisory transitions when desired access differs, and remember the new access.",
                        "The graph does not execute arbitrary runtime behavior by itself. Callback-owned passes hold runtime-specific state; the graph passes them a small context with graph vocabulary and pass identity.",
                        "For DX12, live transition and UAV barriers remain explicit backend-owned calls. If an explicit DX12 helper call does not emit exactly one concrete barrier when it should, the backend fails fatally rather than silently continuing.",
                    ],
                    figure="render_graph",
                    equation="eq_graph_transition",
                ),
                SectionSpec(
                    "Descriptors, Uploads, and Fences",
                    [
                        "DX12 descriptor rows are binding records, not textures. Persistent texture descriptors are copied into transient shader-visible rows before draw so the command list sees a valid table for the current frame.",
                        "Constant buffers are uploaded through aligned upload memory. ShaderDX12 reflects compiled HLSL constant buffers so SetFloat, SetVec, and SetMat writes land at the offsets the shader actually compiled with.",
                        "Present closes the command list, submits it, flips the swapchain, signals a fence, advances allocator/backbuffer indices, waits only when needed before reusing allocator/upload/descriptor memory, and retires deferred releases after completion.",
                    ],
                    figure="dx12_memory",
                    equation="eq_fence",
                ),
                SectionSpec(
                    "Shader Contracts",
                    [
                        "Shader source assets resolve logical names. ShaderDX12 compiles vertex and pixel entry points from the HLSL file, hashes bytecode, reflects uniforms/resources, and checks debug contracts when available.",
                        "ShaderContracts.h documents the current binding ABI: ordinary raster resources use SRV register slots, with slots t0 through t4 currently exposed. t4 is the object material table while per-instance streams carry draw-local material payload.",
                        "Contracts matter because render bugs often come from mismatched assumptions: a shader expects a uniform or texture that the pass never sets, or a pass binds a texture into the wrong slot.",
                    ],
                    bullets=[
                        "Reflection discovers constant-buffer variable offsets and sizes.",
                        "Constant buffer size is aligned to 256 bytes for DX12.",
                        "Debug builds report missing required uniforms and reflection mismatches.",
                    ],
                ),
                SectionSpec(
                    "Cinematic and Debug Rendering",
                    [
                        "Cinematic rendering wraps the frame with HDR targets, sky atmosphere, clouds, god rays, volumetric light, bloom, fog, tone mapping, shadows, and terrain relief settings. These are presentation controls layered over scene state.",
                        "Debug rendering is intentionally part of the render pass order. Collision state colors, broadphase cells, contact manifolds, sleep islands, terrain probes, transparent bodies, replay paths, and UI panels all need to draw in predictable layers.",
                        "The manual distinction is simple: cinematic controls change how the world is seen; physics debug controls show how the world was simulated. Mixing those responsibilities would undermine both visual polish and deterministic evidence.",
                    ],
                    bullets=[
                        "Shadow maps work in normal and cinematic rendering.",
                        "Reflection may use DXR or mirrored raster depending on capability and debug alpha needs.",
                        "The tonemap pass is where HDR scene color, fog, bloom, grade, and volumetric light meet.",
                    ],
                ),
            ],
            [
                "Keep moving pass resource ownership from Run-shaped aggregates into renderer/pass-owned boundaries.",
                "Keep render graph callback coverage honest without moving DX12 barrier ownership out of the backend.",
                "Add diagrammed resource-state examples for shadow, reflection, water, and tonemap resources.",
                "Continue DX12 validation gate improvements around InfoQueue errors and screenshot timing.",
                "Preserve future portability vocabulary without adding new runtime dependency on retired backends.",
            ],
        ),
        ChapterSpec(
            5,
            "Diagnostics, Replay, and Validation",
            "Skullbonez treats evidence as part of the engine. Diagnostics, replay, and validation gates exist so changes can be understood, reproduced, and bounded.",
            "validation_map",
            [
                SectionSpec(
                    "Validation Philosophy",
                    [
                        "Repository validation scripts are pre-commit and PR gates, not a reflex after every edit. The right command is the smallest command that proves the risk introduced by the change.",
                        "Documentation-only changes require no repository validation. Physics behavior changes require byte-exact physics gates. Renderer/shader/screenshot changes require DX12 renderer validation. Hot-path work may require perf validation.",
                        "The non-negotiable part is evidence: never claim validation success without command output. If validation is skipped because the work is docs-only, say that plainly.",
                    ],
                    table=(
                        ["Change Area", "Evidence"],
                        [
                            ["Docs/manual prose", "Render and visual QA for the document; no repo validation."],
                            ["Physics/contact/solver", "tools\\validate_physics.bat and byte-exact CSV match."],
                            ["DX12/shader/rendering", "tools\\validate_dx12_renderer.bat and zero DX12 validation errors."],
                            ["Hot path/perf", "tools\\validate_perf.bat plus manual hot-path review."],
                        ],
                    ),
                ),
                SectionSpec(
                    "SkullScope and Physics Queries",
                    [
                        "Physics diagnostics should stay cheap for model analysis. Instead of loading raw CSV, NDJSON, or SQLite artifacts, use SkullScope traces and query them with tools\\physics_query.bat.",
                        "The physics pipeline overlay records bounded pipeline_stages rows. A query can summarize frames, events, contacts, sleep decisions, or replay restore outcomes without dumping an entire trace into a conversation.",
                        "When SkullScope is used, the handoff must report trace command, query commands, output sizes, and what data the model actually ingested. That discipline keeps diagnostics powerful without turning them into unbounded context sludge.",
                    ],
                    bullets=[
                        "Use --physics-diag in Debug builds to generate queryable traces.",
                        "Prefer summary/events/questions queries over raw artifacts.",
                        "Rerun narrower queries if shell output is truncated.",
                    ],
                ),
                SectionSpec(
                    "Replay as a Runtime Microscope",
                    [
                        "Replay capture stores recent presentation and solver samples. The presentation row previews camera/body state. The solver row stores enough authoritative physics state to restore a retained fixed tick.",
                        "Scrubbing is intentionally presentation-safe: render preview applies a historical pose, draws it, and restores live state afterward. Branching from a solver target is stricter because it mutates live simulation and must restore/check solver state.",
                        "Prediction is a sandboxed lookahead. It advances speculative physics in amortized slices, restores the real live world before rendering, and exposes profiler markers so expensive prediction work can be measured.",
                    ],
                    bullets=[
                        "Replay paths show causal body chains for selected targets.",
                        "Velocity edit pauses simulation and redraws future paths from edited live velocity.",
                        "V2 replay artifacts include event/checkpoint data for file-backed scrub and restore probes.",
                    ],
                ),
                SectionSpec(
                    "Profiler and Memory Evidence",
                    [
                        "The engine exposes CPU profiler scopes and optional platform profiler markers. Physics and rendering code names important stages so investigations can compare time against the architecture.",
                        "Graphics stress is opt-in. It churns scene loads, cinematic settings, sky/fog/ray controls, render toggles, debug overlays, object counts, UI state, and memory records from a deterministic seed.",
                        "Memory evidence separates process memory from engine buckets, DXGI usage, descriptors, upload arena pressure, and cache/pool counts. The goal is not just to know memory grew; it is to know which owner grew.",
                    ],
                    bullets=[
                        "Use platform profiler markers when changing marker code.",
                        "Use bounded graphics stress for resource lifetime or memory-growth investigations.",
                        "Report log paths and key lines instead of pasting giant artifacts.",
                    ],
                ),
                SectionSpec(
                    "Documentation as Evidence",
                    [
                        "This manual is itself a documentation artifact, so its validation is layout validation: render the DOCX, inspect pages, and provide the PDF for printing. It does not prove engine behavior.",
                        "Good technical documentation should also expose uncertainty. Each chapter ends with room for improvement because the engine is actively modernizing; a print manual should teach current reality without pretending the work is finished.",
                        "When the code changes, regenerate or revise the manual from the same source boundaries: startup docs, current source, Agentic references, and rendered QA.",
                    ],
                    bullets=[
                        "Keep statements tied to actual code paths and handoff docs.",
                        "Prefer diagrams for ownership and equations for solver math.",
                        "Leave improvement lists as future-work prompts, not vague criticism.",
                    ],
                ),
            ],
            [
                "Add a repeatable manual regeneration target if the manual becomes a maintained artifact.",
                "Create a small gallery of known-good SkullScope query examples for common physics questions.",
                "Turn validation map tables into links inside the developer docs.",
                "Add a print-friendly replay/diagnostics appendix with concrete command recipes.",
            ],
        ),
        ChapterSpec(
            6,
            "Working With the Engine",
            "The best way to change Skullbonez Core is to respect ownership, validate the specific risk, and leave better explanations than you found.",
            None,
            [
                SectionSpec(
                    "Reading a Subsystem",
                    [
                        "Start with the owner and the data shape. For runtime work, ask what Run owns and which extracted controller or context narrows the operation. For physics, ask whether the code runs on model mirrors or store records. For rendering, ask whether the pass owns resources or borrows a frame context.",
                        "Then ask where the behavior becomes observable. Physics becomes observable through body writeback, CSV baselines, SkullScope traces, debug overlays, and replay. Rendering becomes observable through InfoQueue validation, screenshots, backbuffer capture, and present-time fences.",
                        "Finally, find the improvement boundary. If a file contains a transitional facade, understand whether the change narrows ownership or merely renames a bridge.",
                    ],
                    bullets=[
                        "Owner first, behavior second, validation third.",
                        "For hot paths, prefer arrays, records, scratch buffers, and side-effect queues.",
                        "For public surfaces, preserve command-line and scene-file compatibility unless a plan says otherwise.",
                    ],
                ),
                SectionSpec(
                    "Changing Physics Safely",
                    [
                        "Physics is deterministic, so even small ordering changes matter. A refactor that preserves final-looking motion can still change byte-exact baselines if it changes broadphase pair order, warm-start cache keys, row iteration order, sleep timing, or final integration order.",
                        "The safest physics edits keep hot loops data-oriented, avoid new polymorphic service objects, and commit owner mutations after the pass. If a body needs to wake or a debug record needs to be emitted, queue the fact and apply it at the owning boundary.",
                        "When behavior intentionally changes, refresh baselines only from the final Debug executable, committed scene files, and committed config state, then rerun the matching physics gate.",
                    ],
                    bullets=[
                        "Never treat copied physics artifacts as trustworthy baselines.",
                        "Do not introduce callback-style bridges into solver or collision loops.",
                        "Use SkullScope queries when diagnosing before changing code.",
                    ],
                ),
                SectionSpec(
                    "Changing Rendering Safely",
                    [
                        "DX12 renderer work is explicit-state work. Look for resource lifetime, descriptor lifetime, upload allocator reuse, fence ordering, command list state, and graph transitions before changing pass behavior.",
                        "Visual regressions are measured with DX12 screenshots and zero validation errors. GL/DX11 parity is historical; do not add new runtime dependencies on retired backends to prove a current rendering change.",
                        "If a pass needs a resource, decide whether it is a persistent pass-owned resource, a frame-local graph transient, an imported backbuffer/depth target, or a source asset that the backend materializes.",
                    ],
                    bullets=[
                        "Graph access vocabulary should describe resource intent before DX12 state translation.",
                        "Descriptors are not resources; they are views into resources.",
                        "Fence completion is the authority for reuse and deferred release.",
                    ],
                ),
                SectionSpec(
                    "Writing Helpful Comments and Handoffs",
                    [
                        "The repository comment standard treats comments as learning aids, not decoration. A good comment names a concept, invariant, lifetime, hazard, unit, validation-sensitive behavior, or local vocabulary that the code cannot express cleanly by itself.",
                        "Handoffs should record changed areas, validation decisions, commands run, meaningful results, artifacts updated, and residual risk. Commit notes should explain what changed and why, not merely name files.",
                        "The manual follows the same spirit: diagrams teach ownership, equations teach physics/math, and improvement sections name what is still worth doing.",
                    ],
                    bullets=[
                        "Comment dense or risky local behavior close to the code.",
                        "Avoid vague migration nouns unless the owner, reason, deletion condition, and checker budget are recorded.",
                        "Keep final reports concise but evidence-backed.",
                    ],
                ),
            ],
            [
                "Turn this chapter into a shorter onboarding checklist for new contributors.",
                "Add small worked examples for common tasks: add a scene, add an asset, diagnose a contact, add a render pass.",
                "Create a maintained glossary linked from README and Agentic references.",
                "Keep future manual revisions tied to actual source and rendered QA instead of static memory.",
            ],
        ),
    ]


def add_cover(doc: Document, figures: dict[str, Path]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Skullbonez Core")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Technical Manual")
    r.font.size = Pt(24)
    r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Engine architecture, physics, rendering, diagnostics, and improvement notes")
    r.font.size = Pt(12.5)
    r.font.color.rgb = MUTED

    add_figure(doc, figures["architecture_map"], "Figure 0.1 - The manual's top-level map.", width=5.7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared July 2026 from the current SkullbonezCore repository state.")
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = MUTED
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_heading(doc, "Contents", 1)
    add_callout(
        doc,
        "Print recommendation",
        "Use the DOCX as the editable master and the exported PDF as the Officeworks print copy. The PDF fixes pagination, figures, and equation cards for predictable output.",
        LIGHT_BLUE,
    )
    for chapter in chapter_specs():
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"Chapter {chapter.number}: {chapter.title}").bold = True
        p.add_run(f" - {chapter.thesis}")
    add_heading(doc, "How to Read This Manual", 2)
    for item in [
        "Start each chapter with its mental model, then use the diagrams to place the subsystem in the engine.",
        "Read equations as implementation anchors, not as a claim that every line mirrors a textbook derivation.",
        "Use each Room for Improvement section as a prompt list for future plans or PR review questions.",
        "When source and manual disagree, trust the current source and revise the manual.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
    doc.add_page_break()


def add_section(doc: Document, spec: SectionSpec, figures: dict[str, Path]) -> None:
    add_heading(doc, spec.title, 2)
    for idx, paragraph in enumerate(spec.paragraphs):
        add_para(doc, paragraph)
        if idx == 0 and spec.note:
            add_callout(doc, "Reader note", spec.note, LIGHT_GRAY)
    if spec.figure:
        add_figure(doc, figures[spec.figure], f"Figure - {spec.title}.", width=6.2)
    if spec.equation:
        add_figure(doc, figures[spec.equation], f"Equation - {spec.title}.", width=5.9)
    if spec.bullets:
        add_bullets(doc, spec.bullets)
    if spec.table:
        headers, rows = spec.table
        widths = [1.55, 2.35, 2.45] if len(headers) == 3 else [6.35 / len(headers)] * len(headers)
        add_table(doc, headers, rows, widths)


def add_improvements(doc: Document, chapter: ChapterSpec) -> None:
    add_heading(doc, "Room for Improvement", 2)
    add_para(
        doc,
        "The list below is deliberately framed as engineering opportunity rather than criticism. It names the places where the chapter's current architecture can become clearer, smaller, or easier to validate.",
    )
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [0.55, 5.8])
    table.rows[0].cells[0].text = ""
    table.rows[0].cells[1].text = "Improvement prompt"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_borders(cell)
    for item in chapter.improvements:
        cells = table.add_row().cells
        cells[0].text = "[ ]"
        cells[1].text = item
        for cell in cells:
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_chapter(doc: Document, chapter: ChapterSpec, figures: dict[str, Path]) -> None:
    add_heading(doc, f"Chapter {chapter.number}: {chapter.title}", 1)
    add_callout(doc, "Mental model", chapter.thesis, LIGHT_BLUE)
    if chapter.figure:
        add_figure(doc, figures[chapter.figure], f"Figure {chapter.number}.1 - {chapter.title} overview.", width=6.2)
    for spec in chapter.sections:
        doc.add_page_break()
        add_section(doc, spec, figures)
    doc.add_page_break()
    add_improvements(doc, chapter)
    doc.add_page_break()


def add_appendix(doc: Document) -> None:
    add_heading(doc, "Appendix A: Source Landmarks", 1)
    add_para(
        doc,
        "These paths were used as grounding context for the manual. They are listed so a future reader can reopen the current implementation instead of treating the prose as a substitute for source review.",
    )
    headers = ["Area", "Primary paths"]
    rows = [
        ["Startup and rules", "AGENTS.md; README.md; Agentic/README.md; Agentic/SessionState.md"],
        ["Runtime reference", "Agentic/Reference/runtime-reference.md"],
        ["Class structure", "Agentic/Reference/skullbonez-core-class-structure.md"],
        ["Physics overview", "Agentic/Reference/physics-overview.md"],
        ["Runtime frame/render", "SkullbonezSource/Runtime/App/RunFrame.cpp; SkullbonezSource/Runtime/App/RunRender.cpp"],
        ["Physics", "PhysicsEngine.*; PhysicsScene.cpp; PhysicsWorld.cpp; PersistentContactSolver.*; SpatialGrid.*"],
        ["Runtime", "SkullbonezSource/Runtime/Simulation/SimulationSystem.cpp"],
        ["Rendering", "SkullbonezSource/Runtime/Render/*; SkullbonezSource/Rendering/RenderGraph.*; SkullbonezSource/Rendering/DX12/*"],
        ["Scene and assets", "SkullbonezSource/Scene/*; SkullbonezSource/Assets/AssetSystem.*; SkullbonezData/scenes/*; SkullbonezData/assets/*"],
    ]
    add_table(doc, headers, rows, [1.7, 4.65])

    add_heading(doc, "Appendix B: Glossary", 1)
    glossary = [
        ("Broadphase", "Cheap collision phase that finds pairs worth exact testing."),
        ("Middle phase", "Pair pruning between grid candidates and exact narrowphase."),
        ("Narrowphase", "Precise collision/contact generation for a candidate pair."),
        ("Manifold", "Set of contact points, normals, tangents, and metadata for a touching pair."),
        ("Persistent row", "A solver row with stable feature id and cached impulses across frames."),
        ("Warm start", "Applying last frame's converged impulse before this frame's solver iterations."),
        ("Render graph", "API-neutral declaration of render resources, pass uses, transitions, and transient lifetimes."),
        ("Descriptor", "DX12 binding record that tells shaders how to interpret a resource."),
        ("Fence", "GPU timeline value that proves when frame resources can be reused or released."),
        ("SkullScope", "Queryable diagnostics workflow for compact trace inspection."),
    ]
    for term, meaning in glossary:
        p = doc.add_paragraph()
        p.add_run(term + ": ").bold = True
        p.add_run(meaning)

    add_heading(doc, "Appendix C: Equation Index", 1)
    for key, caption in [
        ("eq_fixed_step", "Fixed timestep accumulator."),
        ("eq_euler", "Velocity-first integration."),
        ("eq_contact_velocity", "Contact point velocity."),
        ("eq_effective_mass", "Normal row impulse update."),
        ("eq_friction", "Tangent friction clamp."),
        ("eq_render_transform", "World to clip space."),
        ("eq_graph_transition", "Render graph transition rule."),
        ("eq_fence", "Fence-scoped reuse."),
    ]:
        add_figure(doc, FIGURES / f"{key}.png", caption, width=5.7)

    doc.add_page_break()
    add_heading(doc, "Appendix D: Worked Reading Pages", 1)
    add_para(
        doc,
        "These pages are short field guides for common reading tasks. They are intentionally procedural so a printed copy can sit beside the code while investigating.",
    )
    for guide in worked_guides():
        doc.add_page_break()
        add_section(doc, guide, {name: FIGURES / f"{name}.png" for name in []})


def build_manual() -> None:
    figures = build_figures()
    doc = Document()
    configure_document(doc)
    add_cover(doc, figures)
    add_contents(doc)
    for chapter in chapter_specs():
        add_chapter(doc, chapter, figures)
    add_appendix(doc)

    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE

    doc.save(OUT_DOCX)
    build_pdf(figures)


def worked_guides() -> list[SectionSpec]:
    return [
        SectionSpec(
            "Trace One Fixed Physics Tick",
            [
                "Start at SimulationSystem::Tick and identify whether the scene is fixed-step or accumulator-driven. Follow the committed tick into GameModelCollection::RunPhysics, then through PhysicsEngine::Step and PhysicsScene::RunPhysics.",
                "In PhysicsScene, watch the model-owner boundary: cold metadata and initial hot fields reload from GameModel state, collider snapshots refresh only when topology changed, PhysicsWorld mutates store-owned hot arrays, then the model-owner mirror writes once for downstream consumers.",
                "In PhysicsWorld, read the phase order as an invariant: per-frame buffers, force pass, broadphase, object CCD front-end, terrain detection, persistent contact solve, point joints, sleep support, remaining integration, sleep island decision.",
            ],
            bullets=[
                "Question to answer: which phase first changes velocity?",
                "Question to answer: which phase first changes position?",
                "Question to answer: which phase writes solved store state back to models?",
                "Validation thought: any phase reorder can change byte-exact physics baselines.",
            ],
        ),
        SectionSpec(
            "Inspect a Contact Row",
            [
                "A contact row is a compact constraint record. It is not a force object and it is not a callback into the model. Read bodyA/bodyB, featureId, normal, tangent axes, rA/rB arms, penetration, support flags, and accumulated impulses.",
                "Then read precompute. Effective mass translates a desired change in contact velocity into an impulse. Bias represents restitution or Baumgarte separation. The friction limit decides how much tangent impulse can be retained.",
                "Finally read the iteration loop. The row computes relative velocity, updates accumulated normal lambda, clamps it to zero or above, updates tangent lambdas, clamps their vector length, and applies only the delta impulse to scratch velocities.",
            ],
            bullets=[
                "Normal row: push apart only.",
                "Tangent rows: resist sliding within a friction budget.",
                "Warm start: previous lambda is applied before iteration.",
                "Cache store: stable feature ids make the next tick start near convergence.",
            ],
        ),
        SectionSpec(
            "Follow a Render Pass",
            [
                "Start from Run::Render. Confirm replay preview state is applied only for the draw and restored afterward. Then follow FrameEntryContext into BuildRenderFrameContext to see which scene values cross the frame boundary and which resources remain owned by RuntimeRenderer.",
                "Inside RuntimeRenderer, find whether the pass is callback-owned through RenderGraph or still executes through a direct pass method. The difference affects where resources are declared, where barriers are emitted, and which diagnostics represent the pass.",
                "For a DX12 issue, keep two maps in your head: the logical render graph resource state and the concrete ID3D12Resource state. The graph should explain intent; RenderBackendDX12 should translate that intent into exactly the barriers DX12 needs.",
            ],
            bullets=[
                "Question to answer: does the pass own persistent resources, graph transients, or imported targets?",
                "Question to answer: are descriptors persistent source rows or frame-visible copies?",
                "Validation thought: renderer changes need zero DX12 validation errors plus screenshot evidence.",
            ],
        ),
        SectionSpec(
            "Add a Reusable Scene Asset",
            [
                "Reusable placeables belong in asset library data, not in one-off editor-only hardcoding. Start with SkullbonezData/assets and decide whether the asset is a convex hull, compound stack, prop, building, or terrain dressing.",
                "Register the asset library from Run setup with an assetlib.* logical name, then reference it from scene assetInstances. If the asset includes hull data, bake hull metadata so source vertices, faces, edges, mass, and inertia stay current.",
                "When testing, separate scene load evidence from physics/render evidence. A new asset can be a parser/scene change, a physics hull change, a render material change, or all three.",
            ],
            bullets=[
                "Prefer assetInstances[] for reusable props.",
                "Bake changed hulls with tools\\bake_hulls.py --write.",
                "Scene or asset data changes are not documentation-only; choose the matching validation gate before PR prep.",
            ],
        ),
        SectionSpec(
            "Read a Scene File",
            [
                "Treat scene JSON as declarative intent. First identify playback and simulation fields: fixedStep, frame count, exit behavior, seed, time scale, and world overrides. Then read objects, asset libraries, asset instances, cameras, terrain, debug, cinematic, and UI fields.",
                "Next ask which setup path consumes the data. Authored scene setup preserves model insertion order and validation gates. Generated setup preserves deterministic demo and solver object algorithms.",
                "Finally ask whether a field is physics-affecting or render-only. Cinematic relief and styles should not alter physics terrain. Object materials alter presentation and material response only when the parsed field explicitly does so.",
            ],
            bullets=[
                "Bare scene names resolve through SkullbonezData\\scenes.",
                "Style includes should not rebuild models or reset physics.",
                "Snapshot output is a serialization path, not the source of truth for hand-authored scenes.",
            ],
        ),
        SectionSpec(
            "Choose a Validation Gate",
            [
                "Validation starts from risk, not habit. If the change is prose-only, repository validation is not required. If it changes physics behavior, run the physics gate. If it changes DX12 rendering, run the DX12 renderer gate. If it touches a hot path, add perf evidence.",
                "Use broad validation only when the scope is broad or uncertain. During iteration, targeted builds or focused probes are acceptable when they answer a specific implementation question, but formal scripts are PR gates.",
                "Always preserve the command and meaningful result lines in the handoff. A future reader should know what was proven, what was not run, and why.",
            ],
            bullets=[
                "Docs/manual: document render QA, no repo validation.",
                "Physics: tools\\validate_physics.bat.",
                "Renderer/shaders: tools\\validate_dx12_renderer.bat.",
                "Unsure at PR gate: tools\\agent_validate.bat.",
            ],
        ),
        SectionSpec(
            "Debug a Physics Regression",
            [
                "Start by deciding whether the regression is deterministic behavior, visual debug output, sleep policy, broadphase reachability, terrain support, or replay restore. That choice determines which artifact is useful.",
                "Prefer SkullScope query output over raw logs. Generate a deterministic trace with --physics-diag, run tools\\physics_query.bat summary and events, then ask focused frame/body/contact/island questions.",
                "If a baseline changed intentionally, regenerate it only from the final Debug executable and committed scene/config state, then rerun the matching gate so the committed baseline is byte-exact against the committed behavior.",
            ],
            bullets=[
                "Do not paste or ingest whole CSV/NDJSON/SQLite artifacts unless raw logs are explicitly requested.",
                "Report trace command, query commands, output sizes, and whether output was truncated.",
                "Treat a single byte of CSV drift as real behavior until proven intentional.",
            ],
        ),
        SectionSpec(
            "Update This Manual Safely",
            [
                "Regenerate diagrams and equations from the builder when layout or content changes. The DOCX is the editable master, but the print PDF should be regenerated from the same content so the two artifacts do not drift.",
                "Before changing technical claims, re-open the current source or Agentic reference path named in Appendix A. The manual should be a guide to the implementation, not a fossilized memory of a previous branch.",
                "After generation, render or inspect the final PDF pages. If DOCX-to-PDF conversion is available on the machine, compare the native PDF and the DOCX-derived PDF for pagination differences before sending to print.",
            ],
            bullets=[
                "Keep chapter improvement lists actionable.",
                "Add new equations only when they clarify implementation behavior.",
                "Update source landmarks when subsystem ownership moves.",
            ],
        ),
    ]


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ManualTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=30,
            leading=36,
            textColor=colors.HexColor("#0B2545"),
            alignment=TA_CENTER,
            spaceAfter=10,
        ),
        "subtitle": ParagraphStyle(
            "ManualSubtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#1F4D78"),
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
        "h1": ParagraphStyle(
            "ManualH1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=8,
            spaceAfter=10,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "ManualH2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=8,
            spaceAfter=7,
            keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "ManualH3",
            parent=base["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=6,
            spaceAfter=5,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "ManualBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10.6,
            leading=14.0,
            textColor=colors.HexColor("#1A202C"),
            spaceAfter=6,
            alignment=TA_LEFT,
        ),
        "caption": ParagraphStyle(
            "ManualCaption",
            parent=base["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=8.8,
            leading=11,
            textColor=colors.HexColor("#5E6C84"),
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "bullet": ParagraphStyle(
            "ManualBullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10.2,
            leading=13.2,
            textColor=colors.HexColor("#1A202C"),
            leftIndent=0.25 * inch,
            firstLineIndent=-0.12 * inch,
            bulletIndent=0.12 * inch,
            spaceAfter=4,
        ),
        "callout_title": ParagraphStyle(
            "ManualCalloutTitle",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=10.2,
            leading=13,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=3,
        ),
        "callout_body": ParagraphStyle(
            "ManualCalloutBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.8,
            leading=12.8,
            textColor=colors.HexColor("#1A202C"),
            spaceAfter=0,
        ),
        "table": ParagraphStyle(
            "ManualTable",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=11,
            textColor=colors.HexColor("#1A202C"),
        ),
        "table_header": ParagraphStyle(
            "ManualTableHeader",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.8,
            leading=11,
            textColor=colors.HexColor("#0F172A"),
        ),
    }


def pdf_on_page(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8.5)
    canvas.setFillColor(colors.HexColor("#5E6C84"))
    canvas.drawString(inch, 10.48 * inch, "Skullbonez Core Technical Manual")
    canvas.drawRightString(7.5 * inch, 0.52 * inch, f"Page {doc.page}")
    canvas.setStrokeColor(colors.HexColor("#D8DEE9"))
    canvas.setLineWidth(0.5)
    canvas.line(inch, 10.38 * inch, 7.5 * inch, 10.38 * inch)
    canvas.restoreState()


def pdf_p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(text), style)


def pdf_callout(story: list, styles: dict[str, ParagraphStyle], title: str, body: str, fill: colors.Color) -> None:
    data = [
        [pdf_p(title, styles["callout_title"])],
        [pdf_p(body, styles["callout_body"])],
    ]
    table = RLTable(data, colWidths=[6.35 * inch])
    table.setStyle(
        RLTableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), fill),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D8DEE9")),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))


def pdf_figure(story: list, styles: dict[str, ParagraphStyle], path: Path, caption: str, width_in: float = 6.2) -> None:
    with Image.open(path) as img:
        width = width_in * inch
        height = width * img.height / img.width
    story.append(RLImage(str(path), width=width, height=height))
    story.append(pdf_p(caption, styles["caption"]))


def pdf_table(
    story: list,
    styles: dict[str, ParagraphStyle],
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_in: Sequence[float],
) -> None:
    data = [[pdf_p(h, styles["table_header"]) for h in headers]]
    for row in rows:
        data.append([pdf_p(cell, styles["table"]) for cell in row])
    table = RLTable(data, colWidths=[w * inch for w in widths_in], repeatRows=1)
    table.setStyle(
        RLTableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))


def pdf_bullets(story: list, styles: dict[str, ParagraphStyle], items: Iterable[str]) -> None:
    for item in items:
        story.append(Paragraph(escape(item), styles["bullet"], bulletText="-"))


def pdf_section(story: list, styles: dict[str, ParagraphStyle], spec: SectionSpec, figures: dict[str, Path]) -> None:
    story.append(Paragraph(escape(spec.title), styles["h2"]))
    for idx, paragraph in enumerate(spec.paragraphs):
        story.append(pdf_p(paragraph, styles["body"]))
        if idx == 0 and spec.note:
            pdf_callout(story, styles, "Reader note", spec.note, colors.HexColor("#F4F6F9"))
    if spec.figure:
        pdf_figure(story, styles, figures[spec.figure], f"Figure - {spec.title}.")
    if spec.equation:
        pdf_figure(story, styles, figures[spec.equation], f"Equation - {spec.title}.", width_in=5.9)
    if spec.bullets:
        pdf_bullets(story, styles, spec.bullets)
    if spec.table:
        headers, rows = spec.table
        widths = [1.55, 2.35, 2.45] if len(headers) == 3 else [6.35 / len(headers)] * len(headers)
        pdf_table(story, styles, headers, rows, widths)


def build_pdf(figures: dict[str, Path]) -> None:
    OUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.9 * inch,
        bottomMargin=0.85 * inch,
        title="Skullbonez Core Technical Manual",
        author="SkullbonezCore",
        subject="Engine architecture, physics, rendering, diagnostics, and improvement notes",
    )
    story: list = []

    story.append(Spacer(1, 72))
    story.append(Paragraph("Skullbonez Core", styles["title"]))
    story.append(Paragraph("Technical Manual", styles["subtitle"]))
    story.append(
        Paragraph(
            "Engine architecture, physics, rendering, diagnostics, and improvement notes",
            styles["subtitle"],
        )
    )
    pdf_figure(story, styles, figures["architecture_map"], "Figure 0.1 - The manual's top-level map.", width_in=5.7)
    story.append(Spacer(1, 16))
    story.append(
        Paragraph(
            "Prepared July 2026 from the current SkullbonezCore repository state.",
            styles["caption"],
        )
    )
    story.append(PageBreak())

    story.append(Paragraph("Contents", styles["h1"]))
    pdf_callout(
        story,
        styles,
        "Print recommendation",
        "Use the DOCX as the editable master and this PDF as the Officeworks print copy. The PDF fixes pagination, figures, and equation cards for predictable output.",
        colors.HexColor("#E8EEF5"),
    )
    for chapter in chapter_specs():
        story.append(
            Paragraph(
                escape(f"Chapter {chapter.number}: {chapter.title}") + " - " + escape(chapter.thesis),
                styles["body"],
                bulletText=str(chapter.number) + ".",
            )
        )
    story.append(Spacer(1, 10))
    story.append(Paragraph("How to Read This Manual", styles["h2"]))
    pdf_bullets(
        story,
        styles,
        [
            "Start each chapter with its mental model, then use diagrams to place the subsystem in the engine.",
            "Read equations as implementation anchors, not as a claim that every line mirrors a textbook derivation.",
            "Use each Room for Improvement section as a prompt list for future plans or PR review questions.",
            "When source and manual disagree, trust the current source and revise the manual.",
        ],
    )
    story.append(PageBreak())

    for chapter in chapter_specs():
        story.append(Paragraph(escape(f"Chapter {chapter.number}: {chapter.title}"), styles["h1"]))
        pdf_callout(story, styles, "Mental model", chapter.thesis, colors.HexColor("#E8EEF5"))
        if chapter.figure:
            pdf_figure(story, styles, figures[chapter.figure], f"Figure {chapter.number}.1 - {chapter.title} overview.")
        for spec in chapter.sections:
            story.append(PageBreak())
            pdf_section(story, styles, spec, figures)
        story.append(PageBreak())
        story.append(Paragraph("Room for Improvement", styles["h2"]))
        story.append(
            pdf_p(
                "The list below is deliberately framed as engineering opportunity rather than criticism. It names the places where the chapter's current architecture can become clearer, smaller, or easier to validate.",
                styles["body"],
            )
        )
        rows = [["", "Improvement prompt"]] + [["[ ]", item] for item in chapter.improvements]
        pdf_table(story, styles, rows[0], rows[1:], [0.55, 5.8])
        story.append(PageBreak())

    story.append(Paragraph("Appendix A: Source Landmarks", styles["h1"]))
    story.append(
        pdf_p(
            "These paths were used as grounding context for the manual. They are listed so a future reader can reopen the current implementation instead of treating the prose as a substitute for source review.",
            styles["body"],
        )
    )
    pdf_table(
        story,
        styles,
        ["Area", "Primary paths"],
        [
            ["Startup and rules", "AGENTS.md; README.md; Agentic/README.md; Agentic/SessionState.md"],
            ["Runtime reference", "Agentic/Reference/runtime-reference.md"],
            ["Class structure", "Agentic/Reference/skullbonez-core-class-structure.md"],
            ["Physics overview", "Agentic/Reference/physics-overview.md"],
            ["Runtime frame/render", "SkullbonezSource/Runtime/App/RunFrame.cpp; SkullbonezSource/Runtime/App/RunRender.cpp"],
            ["Physics", "PhysicsEngine.*; PhysicsScene.cpp; PhysicsWorld.cpp; PersistentContactSolver.*; SpatialGrid.*"],
            ["Runtime", "SkullbonezSource/Runtime/Simulation/SimulationSystem.cpp"],
            ["Rendering", "SkullbonezSource/Runtime/Render/*; SkullbonezSource/Rendering/RenderGraph.*; SkullbonezSource/Rendering/DX12/*"],
            ["Scene and assets", "SkullbonezSource/Scene/*; SkullbonezSource/Assets/AssetSystem.*; SkullbonezData/scenes/*; SkullbonezData/assets/*"],
        ],
        [1.7, 4.65],
    )
    story.append(PageBreak())

    story.append(Paragraph("Appendix B: Glossary", styles["h1"]))
    for term, meaning in [
        ("Broadphase", "Cheap collision phase that finds pairs worth exact testing."),
        ("Middle phase", "Pair pruning between grid candidates and exact narrowphase."),
        ("Narrowphase", "Precise collision/contact generation for a candidate pair."),
        ("Manifold", "Set of contact points, normals, tangents, and metadata for a touching pair."),
        ("Persistent row", "A solver row with stable feature id and cached impulses across frames."),
        ("Warm start", "Applying last frame's converged impulse before this frame's solver iterations."),
        ("Render graph", "API-neutral declaration of render resources, pass uses, transitions, and transient lifetimes."),
        ("Descriptor", "DX12 binding record that tells shaders how to interpret a resource."),
        ("Fence", "GPU timeline value that proves when frame resources can be reused or released."),
        ("SkullScope", "Queryable diagnostics workflow for compact trace inspection."),
    ]:
        story.append(Paragraph(f"<b>{escape(term)}:</b> {escape(meaning)}", styles["body"]))
    story.append(PageBreak())

    story.append(Paragraph("Appendix C: Equation Index", styles["h1"]))
    for key, caption in [
        ("eq_fixed_step", "Fixed timestep accumulator."),
        ("eq_euler", "Velocity-first integration."),
        ("eq_contact_velocity", "Contact point velocity."),
        ("eq_effective_mass", "Normal row impulse update."),
        ("eq_friction", "Tangent friction clamp."),
        ("eq_render_transform", "World to clip space."),
        ("eq_graph_transition", "Render graph transition rule."),
        ("eq_fence", "Fence-scoped reuse."),
    ]:
        pdf_figure(story, styles, FIGURES / f"{key}.png", caption, width_in=5.7)

    story.append(PageBreak())
    story.append(Paragraph("Appendix D: Worked Reading Pages", styles["h1"]))
    story.append(
        pdf_p(
            "These pages are short field guides for common reading tasks. They are intentionally procedural so a printed copy can sit beside the code while investigating.",
            styles["body"],
        )
    )
    for guide in worked_guides():
        story.append(PageBreak())
        pdf_section(story, styles, guide, figures)

    doc.build(story, onFirstPage=pdf_on_page, onLaterPages=pdf_on_page)
    page_count = len(PdfReader(str(OUT_PDF)).pages)
    print(f"{OUT_PDF} ({page_count} pages)")


if __name__ == "__main__":
    build_manual()
    print(str(OUT_DOCX))
